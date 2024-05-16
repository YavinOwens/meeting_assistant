import os
import assemblyai as aai
import streamlit as st
from langchain.llms import OpenAI
from docx import Document
from docx.shared import Pt

from keys import my_keys # import your api keys here 

st.set_page_config(layout="wide")

# Set API keys
aai.settings.api_key = ""
OPENAI_API_KEY = ""
# Create a transcriber instance
transcriber = aai.Transcriber()

# Streamlit app
st.title("Audio Transcription and Insight Generation")
st.write("Upload an audio file to transcribe and generate insights.")

# Upload audio file
uploaded_file = st.file_uploader("Choose an audio file...", type=["mp3", "wav", "mp4", "m4a"])

# Dummy summarize_text function
def summarize_text(text):
    # Dummy implementation: In practice, replace this with actual summarization logic.
    # For demonstration purposes, we are just truncating the text.
    max_words = 1000
    words = text.split()
    if len(words) > max_words:
        return ' '.join(words[:max_words]) + '...'
    return text

def create_meeting_document(transcript, insights):
    doc = Document()
    doc.add_heading('Meeting Transcript and Insights', 0)
    
    doc.add_heading('Transcript:', level=1)
    p = doc.add_paragraph(transcript)
    p.style.font.size = Pt(12)
    
    doc.add_heading('Generated Insights:', level=1)
    for idx, insight in enumerate(insights, 1):
        doc.add_heading(f'Insight {idx}', level=2)
        doc.add_paragraph(insight, style='BodyText')
    
    doc.save('Meeting_Transcript_and_Insights.docx')

if uploaded_file is not None:
    # Save the uploaded file to disk
    with open("uploaded_audio.mp4", "wb") as f:
        f.write(uploaded_file.getbuffer())
    st.success("File uploaded successfully!")

    # Transcribe the uploaded audio file
    st.write("Transcribing the audio file...")
    transcript = transcriber.transcribe("uploaded_audio.mp4")
    
    # Save the transcript to a text file
    with open("transcript.txt", "w") as file:
        file.write(transcript.text)

    st.write("Transcript:")
    st.write(transcript.text)

    # Read the transcript from the text file
    with open("transcript.txt", "r") as file:
        saved_transcript = file.read()

    # Summarize the transcript if it's too long
    if len(saved_transcript.split()) > 4000:  # Roughly estimating tokens by word count
        st.write("Transcript is too long, summarizing...")
        summarized_transcript = summarize_text(saved_transcript)
    else:
        summarized_transcript = saved_transcript

    st.write("Summarized Transcript:")
    st.write(summarized_transcript)

    # Define the prompt for LangChain
    prompt = f"""
    Professional Transcription: The application should act as a professional note-taker, providing world-class transcriptions that detail both high-level and low-level points.
    Data Analytics: Offer possible solutions through data analytics, leveraging master data management principles and governance.
    Action Items: Extract and list actionable items from the transcription.
    Conversation Insights: Analyze the tone and structure of the conversation, providing insights into these aspects.
    Create a new line under each topic. Use a number for each insight in the list.

    Transcription Text:
    {summarized_transcript}
    """

    # Initialize LangChain with OpenAI
    llm = OpenAI(api_key=OPENAI_API_KEY, temperature=0.7)

    # Generate insights based on the prompt
    response = llm.generate([prompt])

    # Extract generated insights
    insights = [output.text for output in response.generations[0]]

    # Display the generated insights
    st.write("Generated Insights:")
    for idx, insight in enumerate(insights, 1):
        st.write(f"Insight {idx}: {insight}")

    # Save the transcript and insights to a Word document
    create_meeting_document(saved_transcript, insights)
    st.success("Meeting transcript and insights saved to 'Meeting_Transcript_and_Insights.docx'")
